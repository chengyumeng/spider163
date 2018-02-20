#!/usr/bin/env python
# -*- coding: utf-8 -*-

from docx import Document
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from sqlalchemy import desc

from spider163 import settings
from spider163.spider import public as uapi
from spider163.utils import tools
from spider163.utils import pylog
from spider163.utils import pysql
from spider163.spider import comment


def read_playlist_json(id):
    url = uapi.playlist_api.format(id)
    data = tools.curl(url,uapi.header)
    return data


def read_music_data(id):
    url = uapi.mp3_url.format(id)


def read_comment_data(id):
    cmt = comment.Comment()
    return cmt.post(id,1)


def read_lyric_data(id):
    url = uapi.lyric_url.format(id)
    data = tools.curl(url,uapi.header)
    return data


def print_pdf(id):
    data = read_playlist_json(id)
    if data["code"] != 200:
        pylog.print_warn("歌单信息拉取失败！")
        return

    document = Document()
    try:
        document.add_heading(data["result"]["name"], 0)
        tags = document.add_paragraph(" ".join(data["result"]["tags"]))
        desc = document.add_paragraph(data["result"]["description"])
        for m in data["result"]["tracks"]:
            document.add_paragraph().add_run(m["name"]).font.color.theme_color = MSO_THEME_COLOR_INDEX.ACCENT_2

            lyric = read_lyric_data(m["id"])
            document.add_paragraph().add_run(lyric["lrc"]["lyric"]).font.color.theme_color = MSO_THEME_COLOR_INDEX.ACCENT_3

            comments = read_comment_data(m["id"])
            for c in comments["hotComments"]:
                author = document.add_paragraph().add_run(c["user"]["nickname"]).style = 'Emphasis'
                content = document.add_paragraph(c["content"])
    except Exception as e:
        pylog.print_warn(e)

    document.save("{}.docx".format(data["result"]["name"]))
    pylog.print_info("文档 {}.docx 已经生成！".format(data["result"]["name"]))


def print_comment(count):
    session = settings.Session()
    comments = session.query(pysql.Comment163).order_by(
        desc(pysql.Comment163.liked)).limit(count)
    document = Document()
    try:
        document.add_heading("TOP {} 评论".format(count), 0)
        i = 0
        for c in comments:
            i = i + 1
            song = session.query(pysql.Music163).filter(pysql.Music163.song_id == c.song_id)
            pylog.print_info("正在填充第 {} 条评论,歌曲：{}".format(i, song[0].song_name))
            document.add_paragraph().add_run(
                "作者：{}".format(c.author)).font.color.theme_color = MSO_THEME_COLOR_INDEX.ACCENT_2
            document.add_paragraph().add_run(
                "内容：{}".format(c.txt)).font.color.theme_color = MSO_THEME_COLOR_INDEX.ACCENT_2
            document.add_paragraph().add_run(
                "歌曲：《{}》 链接：http://music.163.com/#/song?id={}".format(song[0].song_name, c.song_id)).font.color.theme_color = MSO_THEME_COLOR_INDEX.ACCENT_2
            document.add_paragraph().add_run(
                "赞同：{}".format(c.liked)).font.color.theme_color = MSO_THEME_COLOR_INDEX.ACCENT_2
            document.add_paragraph("")
    except Exception as e:
        pylog.print_warn(e)
    document.save("TOP {} 评论.docx".format(count))
    pylog.print_warn("\n完成文档 TOP {} 评论.docx 的生成！\n".format(count))